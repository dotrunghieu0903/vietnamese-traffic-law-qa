#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Đọc và phân tích file DOCX để bổ sung vào file JSON
"""

import json
import re
import os
from typing import Dict, List, Any
from pathlib import Path

def install_required_packages():
    """Cài đặt các package cần thiết"""
    import subprocess
    import sys
    
    packages = ["python-docx", "chardet"]
    
    for package in packages:
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
            print(f"✅ Đã cài đặt {package}")
        except subprocess.CalledProcessError as e:
            print(f"❌ Lỗi khi cài đặt {package}: {e}")

def read_docx_file(file_path: str) -> str:
    """Đọc nội dung từ file DOCX"""
    try:
        from docx import Document
        doc = Document(file_path)
        
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Đọc cả bảng nếu có
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    
    except ImportError:
        print("❌ Chưa cài đặt python-docx. Đang cài đặt...")
        install_required_packages()
        from docx import Document
        return read_docx_file(file_path)
    
    except Exception as e:
        print(f"❌ Lỗi khi đọc file DOCX: {e}")
        return ""

def extract_articles_from_docx(content: str) -> Dict[str, Any]:
    """Trích xuất các điều từ nội dung DOCX"""
    articles = {}
    
    # Pattern để tìm các điều
    article_pattern = r"Điều\s+(\d+)[\.:]?\s*(.+?)(?=Điều\s+\d+|$)"
    
    matches = re.finditer(article_pattern, content, re.DOTALL | re.IGNORECASE)
    
    for match in matches:
        article_num = int(match.group(1))
        article_content = match.group(2).strip()
        
        # Tìm tiêu đề điều
        title_match = re.search(r"^(.+?)(?:\n|Khoản|\d+\.|[a-z]\))", article_content)
        title = title_match.group(1).strip() if title_match else f"Điều {article_num}"
        
        # Tìm các khoản
        sections = extract_sections_from_article(article_content)
        
        articles[f"dieu_{article_num}"] = {
            "title": title,
            "content": article_content[:200] + "..." if len(article_content) > 200 else article_content,
            "sections": sections
        }
    
    return articles

def extract_sections_from_article(article_content: str) -> List[Dict[str, Any]]:
    """Trích xuất các khoản từ một điều"""
    sections = []
    
    # Pattern để tìm khoản
    section_patterns = [
        r"Khoản\s+(\d+)[\.:]?\s*(.+?)(?=Khoản\s+\d+|$)",
        r"(\d+)\.\s*(.+?)(?=\d+\.|$)",
    ]
    
    for pattern in section_patterns:
        matches = re.finditer(pattern, article_content, re.DOTALL | re.IGNORECASE)
        for match in matches:
            section_num = match.group(1)
            section_content = match.group(2).strip()
            
            # Tìm vi phạm và mức phạt
            violations = extract_violations(section_content)
            fine_info = extract_fine_info(section_content)
            
            section_data = {
                "section": f"Khoản {section_num}",
                "content": section_content[:150] + "..." if len(section_content) > 150 else section_content
            }
            
            if violations:
                section_data["violations"] = violations
            
            if fine_info:
                section_data.update(fine_info)
            
            sections.append(section_data)
        
        if sections:  # Nếu đã tìm thấy sections thì dừng
            break
    
    return sections

def extract_violations(content: str) -> List[str]:
    """Trích xuất danh sách vi phạm từ nội dung"""
    violations = []
    
    # Các pattern phổ biến cho vi phạm
    violation_patterns = [
        r"[a-z]\)\s*(.+?)(?=[a-z]\)|$)",
        r"[-−]\s*(.+?)(?=[-−]|$)",
        r"•\s*(.+?)(?=•|$)",
    ]
    
    for pattern in violation_patterns:
        matches = re.finditer(pattern, content, re.MULTILINE)
        for match in matches:
            violation = match.group(1).strip()
            if len(violation) > 10 and len(violation) < 200:  # Lọc bỏ các đoạn quá ngắn hoặc quá dài
                violations.append(violation)
    
    return violations

def extract_fine_info(content: str) -> Dict[str, Any]:
    """Trích xuất thông tin mức phạt từ nội dung"""
    fine_info = {}
    
    # Pattern để tìm mức phạt
    fine_patterns = [
        r"phạt\s+tiền\s+từ\s+([\d.,]+)\s*(?:đến|[-−])\s*([\d.,]+)\s*(?:đồng|VNĐ)",
        r"mức\s+phạt\s+từ\s+([\d.,]+)\s*(?:đến|[-−])\s*([\d.,]+)\s*(?:đồng|VNĐ)",
        r"([\d.,]+)\s*(?:đến|[-−])\s*([\d.,]+)\s*(?:đồng|VNĐ)"
    ]
    
    for pattern in fine_patterns:
        match = re.search(pattern, content, re.IGNORECASE)
        if match:
            min_fine = match.group(1).replace(".", "").replace(",", "")
            max_fine = match.group(2).replace(".", "").replace(",", "")
            fine_info["fine_range"] = f"{min_fine} - {max_fine} VNĐ"
            break
    
    # Tìm biện pháp bổ sung
    additional_measures = []
    measure_patterns = [
        r"tước\s+quyền\s+sử\s+dụng\s+giấy\s+phép\s+lái\s+xe\s+từ\s+(\d+)\s+đến\s+(\d+)\s+tháng",
        r"tạm\s+giữ\s+phương\s+tiện",
        r"tịch\s+thu\s+phương\s+tiện"
    ]
    
    for pattern in measure_patterns:
        matches = re.finditer(pattern, content, re.IGNORECASE)
        for match in matches:
            additional_measures.append(match.group(0))
    
    if additional_measures:
        fine_info["additional_measures"] = additional_measures
    
    return fine_info

def compare_with_existing_json(docx_articles: Dict[str, Any], json_file_path: str) -> Dict[str, Any]:
    """So sánh dữ liệu từ DOCX với JSON hiện tại"""
    try:
        with open(json_file_path, 'r', encoding='utf-8') as f:
            existing_data = json.load(f)
        
        comparison_result = {
            "docx_articles_count": len(docx_articles),
            "json_articles_count": existing_data.get("document_info", {}).get("total_articles", 0),
            "missing_in_json": [],
            "extra_in_json": [],
            "content_differences": []
        }
        
        # Kiểm tra các điều thiếu trong JSON
        existing_articles = existing_data.get("key_articles", {})
        
        for docx_key in docx_articles.keys():
            if docx_key not in existing_articles:
                comparison_result["missing_in_json"].append(docx_key)
        
        # Kiểm tra các điều thừa trong JSON
        for json_key in existing_articles.keys():
            if json_key not in docx_articles:
                comparison_result["extra_in_json"].append(json_key)
        
        # So sánh nội dung
        for key in set(docx_articles.keys()) & set(existing_articles.keys()):
            docx_title = docx_articles[key].get("title", "")
            json_title = existing_articles[key].get("title", "")
            
            if docx_title.lower().strip() != json_title.lower().strip():
                comparison_result["content_differences"].append({
                    "article": key,
                    "field": "title",
                    "docx_value": docx_title,
                    "json_value": json_title
                })
        
        return comparison_result
    
    except Exception as e:
        print(f"❌ Lỗi khi so sánh với JSON: {e}")
        return {}

def generate_updated_json(docx_articles: Dict[str, Any], original_json_path: str, output_path: str) -> bool:
    """Tạo file JSON đã cập nhật với dữ liệu từ DOCX"""
    try:
        with open(original_json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Cập nhật thông tin tài liệu
        if "document_info" in data:
            data["document_info"]["total_articles"] = len(docx_articles)
            data["document_info"]["description"] = f"Nghị định 100/2019 - Đầy đủ {len(docx_articles)} điều được trích xuất từ file DOCX"
        
        # Cập nhật key_articles
        if "key_articles" not in data:
            data["key_articles"] = {}
        
        data["key_articles"].update(docx_articles)
        
        # Cập nhật thống kê
        if "statistics" in data:
            data["statistics"]["total_articles"] = len(docx_articles)
            data["statistics"]["articles_with_violations"] = len([a for a in docx_articles.values() if "sections" in a and any("violations" in s for s in a["sections"])])
        
        # Lưu file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        return True
    
    except Exception as e:
        print(f"❌ Lỗi khi tạo file JSON mới: {e}")
        return False

def main():
    """Hàm chính"""
    # Đường dẫn files
    docx_file = r"c:\Users\hieudt22\Documents\VNI-TrafficLawQA\docs\ND100.docx"
    json_file = r"c:\Users\hieudt22\Documents\VNI-TrafficLawQA\data\raw\legal_documents\nghi_dinh_100_2019.json"
    output_file = r"c:\Users\hieudt22\Documents\VNI-TrafficLawQA\data\processed\nghi_dinh_100_2019_updated.json"
    
    print("🔍 Đang phân tích file DOCX...")
    
    # Kiểm tra file tồn tại
    if not os.path.exists(docx_file):
        print(f"❌ Không tìm thấy file DOCX: {docx_file}")
        return
    
    if not os.path.exists(json_file):
        print(f"❌ Không tìm thấy file JSON: {json_file}")
        return
    
    # Đọc file DOCX
    print("📖 Đang đọc file DOCX...")
    docx_content = read_docx_file(docx_file)
    
    if not docx_content:
        print("❌ Không thể đọc nội dung từ file DOCX")
        return
    
    print(f"✅ Đã đọc {len(docx_content)} ký tự từ file DOCX")
    
    # Trích xuất các điều
    print("🔍 Đang trích xuất các điều từ DOCX...")
    docx_articles = extract_articles_from_docx(docx_content)
    
    print(f"✅ Đã trích xuất {len(docx_articles)} điều từ file DOCX")
    
    # So sánh với JSON hiện tại
    print("🔍 Đang so sánh với file JSON hiện tại...")
    comparison = compare_with_existing_json(docx_articles, json_file)
    
    if comparison:
        print("\n📊 KẾT QUẢ SO SÁNH:")
        print(f"   - Số điều trong DOCX: {comparison['docx_articles_count']}")
        print(f"   - Số điều trong JSON: {comparison['json_articles_count']}")
        print(f"   - Thiếu trong JSON: {len(comparison['missing_in_json'])} điều")
        print(f"   - Thừa trong JSON: {len(comparison['extra_in_json'])} điều")
        print(f"   - Khác biệt nội dung: {len(comparison['content_differences'])} điều")
        
        if comparison['missing_in_json']:
            print(f"\n📝 Các điều thiếu trong JSON: {', '.join(comparison['missing_in_json'])}")
        
        if comparison['extra_in_json']:
            print(f"\n📝 Các điều thừa trong JSON: {', '.join(comparison['extra_in_json'])}")
        
        if comparison['content_differences']:
            print(f"\n📝 Khác biệt nội dung:")
            for diff in comparison['content_differences'][:5]:  # Hiển thị 5 khác biệt đầu tiên
                print(f"   - {diff['article']}: {diff['field']}")
                print(f"     DOCX: {diff['docx_value'][:50]}...")
                print(f"     JSON: {diff['json_value'][:50]}...")
    
    # Tạo file JSON cập nhật
    print(f"\n💾 Đang tạo file JSON cập nhật...")
    if generate_updated_json(docx_articles, json_file, output_file):
        print(f"✅ Đã tạo file JSON cập nhật: {output_file}")
    else:
        print("❌ Lỗi khi tạo file JSON cập nhật")
    
    # Hiển thị một số điều mẫu
    print(f"\n📄 MỘT SỐ ĐIỀU MẪU TỪ DOCX:")
    sample_articles = list(docx_articles.items())[:3]
    for key, article in sample_articles:
        print(f"\n{key.upper()}:")
        print(f"   Tiêu đề: {article['title']}")
        print(f"   Số khoản: {len(article.get('sections', []))}")
        if article.get('sections'):
            first_section = article['sections'][0]
            print(f"   Khoản đầu: {first_section.get('content', '')[:100]}...")

if __name__ == "__main__":
    main()
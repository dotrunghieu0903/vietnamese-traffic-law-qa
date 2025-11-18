#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to extract and parse Vietnamese traffic law content from DOCX file
Specifically designed for ND100-2019 document structure
"""

from docx import Document
import json
import re
from datetime import datetime

def extract_docx_content(docx_path):
    """Extract text content from DOCX file"""
    try:
        doc = Document(docx_path)
        
        # Extract all paragraphs
        content = []
        for para in doc.paragraphs:
            if para.text.strip():  # Only non-empty paragraphs
                content.append(para.text.strip())
        
        # Extract tables if any
        tables_content = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables_content.append(table_data)
        
        return {
            "text_content": content,
            "tables": tables_content
        }
    
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        return None

def parse_vietnamese_traffic_law(content):
    """Parse Vietnamese traffic law content into structured format"""
    if not content or not content.get("text_content"):
        return None
    
    text_lines = content["text_content"]
    parsed_articles = {}
    current_article = None
    current_section = None
    current_violations = []
    current_fine_range = None
    current_additional_measures = []
    
    # Patterns for matching different elements
    article_pattern = r"Điều\s+(\d+)\.\s*(.*)"
    section_pattern = r"(\d+)\.\s*Phạt tiền từ\s*([\d.,]+)\s*đồng đến\s*([\d.,]+)\s*đồng.*:"
    violation_pattern = r"^([a-z]|đ)\)\s*(.*)"
    additional_measure_pattern = r"(\d+)\.\s*Ngoài.*phạt.*bổ sung.*:"
    
    i = 0
    while i < len(text_lines):
        line = text_lines[i].strip()
        
        # Check for article header
        article_match = re.match(article_pattern, line)
        if article_match:
            # Save previous article if exists
            if current_article and current_section:
                if current_article not in parsed_articles:
                    parsed_articles[current_article] = {"sections": []}
                
                section_data = {
                    "section": f"Khoản {current_section}",
                    "fine_range": current_fine_range,
                    "violations": current_violations.copy()
                }
                if current_additional_measures:
                    section_data["additional_measures"] = current_additional_measures.copy()
                
                parsed_articles[current_article]["sections"].append(section_data)
            
            # Start new article
            article_num = article_match.group(1)
            article_title = article_match.group(2)
            current_article = f"dieu_{article_num}"
            
            if current_article not in parsed_articles:
                parsed_articles[current_article] = {
                    "title": article_title,
                    "sections": []
                }
            
            current_section = None
            current_violations = []
            current_fine_range = None
            current_additional_measures = []
            
        # Check for section with fine range
        elif re.match(section_pattern, line):
            # Save previous section if exists
            if current_article and current_section:
                section_data = {
                    "section": f"Khoản {current_section}",
                    "fine_range": current_fine_range,
                    "violations": current_violations.copy()
                }
                if current_additional_measures:
                    section_data["additional_measures"] = current_additional_measures.copy()
                
                parsed_articles[current_article]["sections"].append(section_data)
            
            # Start new section
            section_match = re.match(section_pattern, line)
            current_section = section_match.group(1)
            min_fine = section_match.group(2).replace(".", "")
            max_fine = section_match.group(3).replace(".", "")
            current_fine_range = f"{min_fine} - {max_fine} VNĐ"
            current_violations = []
            current_additional_measures = []
            
        # Check for violation items
        elif re.match(violation_pattern, line):
            violation_match = re.match(violation_pattern, line)
            violation_text = violation_match.group(2)
            
            # Clean up violation text
            violation_text = re.sub(r',\s*trừ.*?;', '', violation_text)
            violation_text = re.sub(r';\s*$', '', violation_text)
            
            # Split multiple violations if separated by semicolon
            if ';' in violation_text:
                violations = [v.strip() for v in violation_text.split(';') if v.strip()]
                current_violations.extend(violations)
            else:
                current_violations.append(violation_text)
        
        # Check for additional measures
        elif "tước quyền" in line.lower() or "tịch thu" in line.lower() or "buộc" in line.lower():
            if "tước quyền sử dụng giấy phép lái xe" in line.lower():
                measure_match = re.search(r"từ\s+(\d+)\s+.*?đến\s+(\d+)\s+tháng", line)
                if measure_match:
                    from_months = measure_match.group(1)
                    to_months = measure_match.group(2)
                    current_additional_measures.append(f"Tước quyền sử dụng giấy phép lái xe từ {from_months} đến {to_months} tháng")
            elif "tịch thu" in line.lower():
                current_additional_measures.append("Tịch thu phương tiện")
            elif "buộc" in line.lower():
                current_additional_measures.append("Buộc khôi phục lại tình trạng ban đầu")
        
        i += 1
    
    # Save the last article and section
    if current_article and current_section:
        section_data = {
            "section": f"Khoản {current_section}",
            "fine_range": current_fine_range,
            "violations": current_violations.copy()
        }
        if current_additional_measures:
            section_data["additional_measures"] = current_additional_measures.copy()
        
        parsed_articles[current_article]["sections"].append(section_data)
    
    return parsed_articles

def format_currency(amount_str):
    """Format currency string for consistency"""
    # Remove dots and convert to standard format
    amount_str = amount_str.replace(".", "")
    return f"{amount_str:,}".replace(",", ".")

def enhance_article_data(parsed_articles, original_content):
    """Enhance parsed articles with additional metadata and formatting"""
    enhanced = {}
    
    for article_key, article_data in parsed_articles.items():
        enhanced[article_key] = {
            "title": article_data["title"],
            "sections": []
        }
        
        # Add source information
        enhanced[article_key]["source_document"] = "ND100-2019.docx"
        enhanced[article_key]["extraction_date"] = datetime.now().strftime("%Y-%m-%d")
        
        # Process sections
        for section in article_data["sections"]:
            enhanced_section = {
                "section": section["section"],
                "fine_range": section["fine_range"],
                "violations": section["violations"]
            }
            
            if "additional_measures" in section and section["additional_measures"]:
                enhanced_section["additional_measures"] = section["additional_measures"]
            
            enhanced[article_key]["sections"].append(enhanced_section)
    
    return enhanced

if __name__ == "__main__":
    docx_path = r"c:\Users\Mr Hieu\Documents\vietnamese-traffic-law-qa\docs\ND100-2019.docx"
    raw_output_path = r"c:\Users\Mr Hieu\Documents\vietnamese-traffic-law-qa\docs\ND100-2019-extracted.json"
    parsed_output_path = r"c:\Users\Mr Hieu\Documents\vietnamese-traffic-law-qa\docs\ND100-2019-parsed-articles.json"
    
    print("Extracting content from DOCX file...")
    content = extract_docx_content(docx_path)
    
    if content:
        # Save raw extracted content
        with open(raw_output_path, 'w', encoding='utf-8') as f:
            json.dump(content, f, ensure_ascii=False, indent=2)
        print(f"Raw content extracted to: {raw_output_path}")
        
        # Parse the content into structured articles
        print("\nParsing articles from extracted content...")
        parsed_articles = parse_vietnamese_traffic_law(content)
        
        if parsed_articles:
            # Enhance the parsed data
            enhanced_articles = enhance_article_data(parsed_articles, content)
            
            # Save parsed articles
            with open(parsed_output_path, 'w', encoding='utf-8') as f:
                json.dump(enhanced_articles, f, ensure_ascii=False, indent=2)
            print(f"Parsed articles saved to: {parsed_output_path}")
            
            # Print summary
            print(f"\n=== PARSING SUMMARY ===")
            print(f"Total articles parsed: {len(enhanced_articles)}")
            
            for article_key, article_data in enhanced_articles.items():
                article_num = article_key.replace("dieu_", "")
                sections_count = len(article_data["sections"])
                total_violations = sum(len(section["violations"]) for section in article_data["sections"])
                print(f"Điều {article_num}: {sections_count} khoản, {total_violations} vi phạm - {article_data['title']}")
        else:
            print("Failed to parse articles from extracted content")
            
        # Also print first few lines for verification
        print("\n=== FIRST 20 LINES FOR VERIFICATION ===")
        for i, para in enumerate(content["text_content"][:20], 1):
            print(f"{i:2d}. {para}")
    else:
        print("Failed to extract content from DOCX file")